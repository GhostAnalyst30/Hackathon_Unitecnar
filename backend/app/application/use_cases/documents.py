"""Casos de uso de documentos: subir, listar, editar, decidir y exportar.

Los routers de la capa api son delgados y delegan aquí; las respuestas se
mapean a DTOs en la capa api.
"""

import io
import uuid
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from markdownify import markdownify
from sqlalchemy import select
from sqlalchemy.orm import Session

from ...config import UPLOADS_DIR
from ...domain.entities import Document
from ...domain.services import anchor_findings
from ...infrastructure.ingest import detect_format, html_to_text
from ...infrastructure.llm import LLMNotConfigured, resolve_chat_config
from .analysis import enqueue_analysis
from .settings import get_settings


class DocumentNotFound(Exception):
    pass


class InvalidOperation(Exception):
    pass


def require_llm_configured(session: Session) -> None:
    """Lanza LLMNotConfigured si no hay proveedor listo (sin modo demo)."""
    resolve_chat_config(get_settings(session))


def get_document(session: Session, document_id: str) -> Document:
    doc = session.get(Document, document_id)
    if doc is None:
        raise DocumentNotFound("Documento no encontrado")
    return doc


def list_documents(session: Session) -> list[Document]:
    return list(
        session.scalars(select(Document).order_by(Document.created_at.desc())).all()
    )


def create_documents(
    session: Session, files: list[tuple[str, bytes]]
) -> list[Document]:
    """Guarda los archivos, crea los documentos y encola su análisis."""
    require_llm_configured(session)
    created: list[Document] = []
    for filename, content in files:
        file_format = detect_format(filename)  # ValueError si no está soportado
        ext = Path(filename).suffix.lower()
        stored_path = UPLOADS_DIR / f"{uuid.uuid4().hex}{ext}"
        stored_path.write_bytes(content)
        doc = Document(
            filename=filename,
            file_path=str(stored_path),
            file_format=file_format,
            status="queued",
        )
        session.add(doc)
        created.append(doc)

    session.commit()
    for doc in created:
        enqueue_analysis(doc.id)
    return created


def delete_document(session: Session, document_id: str) -> None:
    doc = get_document(session, document_id)
    try:
        Path(doc.file_path).unlink(missing_ok=True)
    except OSError:
        pass
    session.delete(doc)
    session.commit()


def update_content(session: Session, document_id: str, content_html: str):
    """Guarda la edición y re-ancla los hallazgos sobre el texto nuevo."""
    doc = get_document(session, document_id)
    doc.content_html = content_html
    doc.content_text = html_to_text(content_html)
    anchor_findings(doc.content_text, doc.findings)
    session.commit()
    return doc.findings


def reanalyze_document(session: Session, document_id: str) -> Document:
    require_llm_configured(session)
    doc = get_document(session, document_id)
    if doc.status in {"queued", "extracting", "analyzing"}:
        raise InvalidOperation("El documento ya se está analizando")
    if not doc.content_text.strip():
        raise InvalidOperation("El documento no tiene contenido para analizar")
    doc.status = "queued"
    session.commit()
    enqueue_analysis(document_id, skip_ingest=True)
    return doc


def decide_document(
    session: Session, document_id: str, decision: str, comment: str | None
) -> Document:
    """Decisión final del humano: validar o descartar."""
    if decision not in {"validated", "discarded"}:
        raise InvalidOperation("La decisión debe ser 'validated' o 'discarded'")
    doc = get_document(session, document_id)
    if doc.status not in {"awaiting_review", "validated", "discarded"}:
        raise InvalidOperation("El documento aún no está listo para revisión")
    doc.status = decision
    doc.decision_comment = comment
    session.commit()
    return doc


def _html_to_docx_bytes(html: str, title: str) -> bytes:
    docx = DocxDocument()
    docx.core_properties.title = title
    soup = BeautifulSoup(html, "html.parser")
    heading_levels = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
    elements = soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote"])
    if not elements:
        docx.add_paragraph(soup.get_text(separator="\n"))
    for el in elements:
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if el.name in heading_levels:
            docx.add_heading(text, level=heading_levels[el.name])
        elif el.name == "li":
            docx.add_paragraph(text, style="List Bullet")
        else:
            docx.add_paragraph(text)
    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


def export_document(
    session: Session, document_id: str, export_format: str
) -> tuple[bytes, str, str]:
    """Devuelve (contenido, media_type, nombre_de_archivo) del texto editado."""
    doc = get_document(session, document_id)
    base_name = Path(doc.filename).stem or "documento"

    if export_format == "md":
        content = markdownify(doc.content_html, heading_style="ATX")
        return content.encode("utf-8"), "text/markdown; charset=utf-8", f"{base_name}.md"
    if export_format == "html":
        return doc.content_html.encode("utf-8"), "text/html; charset=utf-8", f"{base_name}.html"
    if export_format == "docx":
        data = _html_to_docx_bytes(doc.content_html, base_name)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return data, media, f"{base_name}.docx"
    raise InvalidOperation("Formato de exportación no soportado (usa md, html o docx)")
